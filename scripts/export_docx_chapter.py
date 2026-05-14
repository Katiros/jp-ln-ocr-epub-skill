#!/usr/bin/env python3
"""Export reviewed OCR/translation Markdown into DOCX files."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document  # type: ignore
from docx.enum.text import WD_BREAK  # type: ignore
from docx.shared import Pt  # type: ignore


PAGE_MARKER = re.compile(r"^<!--\s*page\s+(\d+)\s*-->$")
HEADING = re.compile(r"^(#{1,3})\s+(.+)$")
QUOTE = re.compile(r"^>\s?(.*)$")


def read_markdown(path: Path | None) -> list[str]:
    if path is None or not path.exists():
        return []
    return path.read_text(encoding="utf-8").splitlines()


def ensure_style(document: Document, name: str, font_name: str, size: int) -> None:
    styles = document.styles
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, 1)
    style.font.name = font_name
    style.font.size = Pt(size)


def setup_styles(document: Document) -> None:
    ensure_style(document, "JP Original", "Yu Mincho", 10)
    ensure_style(document, "CN Translation", "Microsoft YaHei", 10)
    ensure_style(document, "OCR Warning", "Microsoft YaHei", 9)
    ensure_style(document, "Illustration Ref", "Microsoft YaHei", 9)


def add_markdown(document: Document, lines: list[str], paragraph_style: str, include_page_breaks: bool) -> None:
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        page_match = PAGE_MARKER.match(line)
        if page_match:
            paragraph = document.add_paragraph(f"[page {page_match.group(1)}]", style="OCR Warning")
            if include_page_breaks:
                paragraph.runs[0].add_break(WD_BREAK.PAGE)
            continue
        heading_match = HEADING.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            document.add_heading(heading_match.group(2), level=level)
            continue
        quote_match = QUOTE.match(line)
        if quote_match:
            document.add_paragraph(quote_match.group(1), style="OCR Warning")
            continue
        document.add_paragraph(line, style=paragraph_style)


def export_single(title: str, lines: list[str], output: Path, style: str, include_page_breaks: bool) -> None:
    document = Document()
    setup_styles(document)
    document.add_heading(title, level=1)
    add_markdown(document, lines, style, include_page_breaks)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def export_bilingual(title: str, jp_lines: list[str], zh_lines: list[str], output: Path) -> None:
    document = Document()
    setup_styles(document)
    document.add_heading(title, level=1)
    document.add_heading("OCR Japanese", level=2)
    add_markdown(document, jp_lines, "JP Original", include_page_breaks=True)
    if zh_lines:
        document.add_page_break()
        document.add_heading("Chinese Translation", level=2)
        add_markdown(document, zh_lines, "CN Translation", include_page_breaks=False)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Export chapter Markdown to DOCX")
    parser.add_argument("--title", default="Chapter")
    parser.add_argument("--jp", required=True, help="Japanese OCR/review Markdown")
    parser.add_argument("--zh", help="Chinese translation Markdown")
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--prefix", default="ch001")
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    jp_lines = read_markdown(Path(args.jp))
    zh_lines = read_markdown(Path(args.zh) if args.zh else None)

    export_single(args.title + " OCR", jp_lines, output_dir / f"{args.prefix}_ocr.docx", "JP Original", True)
    if zh_lines:
        export_single(args.title + " Chinese", zh_lines, output_dir / f"{args.prefix}_zh.docx", "CN Translation", False)
    export_bilingual(args.title + " Bilingual", jp_lines, zh_lines, output_dir / f"{args.prefix}_bilingual.docx")


if __name__ == "__main__":
    main()

